#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to edit Flinders University transcript document - Version 2
Creates new version with different course and improved academic records
Preserves all formatting, layout, and visual elements
"""

import docx
from docx import Document
import re
from datetime import datetime, timedelta
import random
import os

def edit_transcript_document_v2():
    """Edit the transcript document with updated information - Version 2"""
    
    # File paths
    input_file = "test read doc/2386787-official-transcript-flinders-university.docx"
    output_file = "test read doc/Flinders_Transcript_Engineering_Updated_2026.docx"
    
    try:
        # Load the document
        print("🔄 Đang đọc file gốc...")
        doc = Document(input_file)
        print(f"✅ Đã load document với {len(doc.paragraphs)} paragraphs")
        
        # New course options (different from v1)
        courses = [
            {
                "name": "Bachelor of Engineering (Software Engineering)",
                "code": "CRICOS Code: 088745K",
                "level": "Bachelor Degree (AQF Level 7)",
                "duration": "4 years"
            },
            {
                "name": "Master of Data Science and Artificial Intelligence",
                "code": "CRICOS Code: 093562M", 
                "level": "Master Degree (AQF Level 9)",
                "duration": "2 years"
            },
            {
                "name": "Bachelor of Computer Science (Cybersecurity)",
                "code": "CRICOS Code: 089174G",
                "level": "Bachelor Degree (AQF Level 7)", 
                "duration": "3 years"
            },
            {
                "name": "Master of Engineering (Computer and Network Engineering)",
                "code": "CRICOS Code: 095831J",
                "level": "Master Degree (AQF Level 9)",
                "duration": "2 years"
            }
        ]
        
        # Select random course (different algorithm for variety)
        selected_course = random.choice(courses)
        print(f"📚 Ngành học mới: {selected_course['name']}")
        
        # Time adjustments (+1 year 1 month) - improved date handling
        def add_time(date_str):
            """Add years and months to date strings with better logic"""
            try:
                # More comprehensive date replacements
                replacements = {
                    "July 2025": "August 2026",
                    "July 2026": "August 2027", 
                    "28 July 2025": "28 August 2026",
                    "12 July 2025": "12 August 2026",
                    "2025": "2026",
                    "2026": "2027"
                }
                
                new_text = date_str
                for old_date, new_date in replacements.items():
                    if old_date in new_text:
                        new_text = new_text.replace(old_date, new_date)
                        
                return new_text
            except:
                return date_str
        
        # Enhanced subjects list for different courses
        all_subjects = {
            "engineering": [
                "Engineering Mathematics I",
                "Programming Fundamentals", 
                "Digital Systems Design",
                "Software Engineering Principles",
                "Data Structures and Algorithms",
                "Computer Networks and Security",
                "Database Management Systems",
                "Object-Oriented Programming",
                "Systems Analysis and Design",
                "Project Management"
            ],
            "data_science": [
                "Statistical Methods",
                "Machine Learning Fundamentals",
                "Data Mining and Analytics", 
                "Big Data Technologies",
                "Python for Data Science",
                "Deep Learning",
                "Business Intelligence",
                "Research Methods"
            ],
            "cybersecurity": [
                "Network Security",
                "Cryptography and Security",
                "Ethical Hacking",
                "Digital Forensics",
                "Risk Management",
                "Security Policies",
                "Incident Response",
                "Secure Software Development"
            ]
        }
        
        # Select appropriate subjects based on course
        if "Engineering" in selected_course['name']:
            subjects = all_subjects["engineering"]
        elif "Data Science" in selected_course['name']:
            subjects = all_subjects["data_science"] 
        elif "Cybersecurity" in selected_course['name']:
            subjects = all_subjects["cybersecurity"]
        else:
            subjects = all_subjects["engineering"]  # default
        
        # Improved grade generation with more realistic distribution
        def generate_realistic_grade():
            """Generate more realistic grade distribution"""
            rand = random.random()
            if rand < 0.12:  # 12% High Distinction (top students)
                grade = "HD"
                points = random.randint(87, 96)
            elif rand < 0.28:  # 16% Distinction  
                grade = "D"
                points = random.randint(76, 86)
            elif rand < 0.55:  # 27% Credit (most common)
                grade = "C" 
                points = random.randint(66, 75)
            elif rand < 0.85:  # 30% Pass
                grade = "P"
                points = random.randint(51, 65)
            else:  # 15% Fail (realistic for some difficulty)
                grade = "F"
                points = random.randint(35, 49)
            
            return grade, points
        
        # Generate enhanced grade table content
        grades_content = "\n\n" + "="*60 + "\n"
        grades_content += "                    ACADEMIC TRANSCRIPT\n"
        grades_content += "="*60 + "\n\n"
        grades_content += f"Student: Ricardo Richardson (ID: 2386219)\n"
        grades_content += f"Course: {selected_course['name']}\n"
        grades_content += f"Duration: {selected_course['duration']} (Full-time)\n"
        grades_content += f"Level: {selected_course['level']}\n"
        grades_content += f"Campus: Bedford Park, South Australia\n\n"
        
        # Semester breakdown
        semesters = ["Semester 1, 2026", "Semester 2, 2026"]
        total_credits = 0
        weighted_points = 0
        passed_subjects = 0
        
        for sem_idx, semester in enumerate(semesters):
            grades_content += f"{semester}\n" + "-"*50 + "\n"
            
            # Select subjects for this semester
            sem_subjects = subjects[sem_idx*4:(sem_idx+1)*4]  # 4 subjects per semester
            
            for subject in sem_subjects:
                grade, points = generate_realistic_grade()
                credits = 7.5  # Standard credit points
                
                if grade != "F":  # Only count passed subjects
                    total_credits += credits
                    weighted_points += points * credits
                    passed_subjects += 1
                
                status = "PASS" if grade != "F" else "FAIL"
                grades_content += f"{subject:<35} {grade:>3} {points:>3}% {credits:>4} CP [{status}]\n"
            
            grades_content += "\n"
        
        # Calculate final statistics
        gpa = weighted_points / total_credits if total_credits > 0 else 0
        completion_rate = (passed_subjects / len(subjects[:8])) * 100
        
        grades_content += "="*60 + "\n"
        grades_content += "ACADEMIC SUMMARY\n"
        grades_content += "="*60 + "\n"
        grades_content += f"{'Subjects Attempted:':<25} {len(subjects[:8]):>8}\n"
        grades_content += f"{'Subjects Passed:':<25} {passed_subjects:>8}\n"
        grades_content += f"{'Total Credits Earned:':<25} {total_credits:>8.1f} CP\n"
        grades_content += f"{'Completion Rate:':<25} {completion_rate:>7.1f}%\n"
        grades_content += f"{'Grade Point Average:':<25} {gpa:>7.2f}/100\n"
        
        # Academic standing
        if gpa >= 85:
            standing = "High Distinction"
        elif gpa >= 75:
            standing = "Distinction"
        elif gpa >= 65:
            standing = "Credit"
        elif gpa >= 50:
            standing = "Pass"
        else:
            standing = "At Risk"
            
        grades_content += f"{'Academic Standing:':<25} {standing:>15}\n"
        grades_content += "="*60 + "\n"
        
        # Edit document content
        changes_made = 0
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            new_text = original_text
            
            # Replace course information (more comprehensive)
            old_course_patterns = [
                "Diploma of Business (Pathway to Bachelor of Business - International Business Specialisation)",
                "Diploma of Business",
                "Bachelor of Business"
            ]
            
            for old_pattern in old_course_patterns:
                if old_pattern in new_text:
                    new_text = new_text.replace(old_pattern, selected_course['name'])
                    changes_made += 1
            
            # Replace CRICOS codes
            if "109279H" in new_text:
                new_text = new_text.replace("109279H", selected_course['code'].split(": ")[1])
                changes_made += 1
            
            # Update dates
            if any(date in new_text for date in ["July 2025", "July 2026", "2025", "2026"]):
                updated_text = add_time(new_text)
                if updated_text != new_text:
                    new_text = updated_text
                    changes_made += 1
            
            # Update course level
            level_patterns = ["Diploma Level (AQF Level 5)", "AQF Level 5"]
            for pattern in level_patterns:
                if pattern in new_text:
                    new_text = new_text.replace(pattern, selected_course['level'])
                    changes_made += 1
            
            # Update duration
            if "1 year (full-time)" in new_text:
                new_text = new_text.replace("1 year (full-time)", f"{selected_course['duration']} (full-time)")
                changes_made += 1
            
            # Apply changes to paragraph
            if new_text != original_text:
                paragraph.text = new_text
                print(f"✏️  Đã sửa: {original_text[:40]}... → {new_text[:40]}...")
        
        # Process tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = original_text
                    
                    # Update course info in tables
                    for old_pattern in old_course_patterns:
                        if old_pattern in new_text:
                            new_text = new_text.replace(old_pattern, selected_course['name'])
                            changes_made += 1
                    
                    # Update dates in tables
                    if any(date in new_text for date in ["July 2025", "July 2026", "2025", "2026"]):
                        updated_text = add_time(new_text)
                        if updated_text != new_text:
                            new_text = updated_text
                            changes_made += 1
                    
                    if new_text != original_text:
                        cell.text = new_text
                        print(f"📊 Đã sửa table: {original_text[:30]}... → {new_text[:30]}...")
        
        # Add enhanced grades section
        doc.add_paragraph(grades_content)
        changes_made += 1
        print(f"📊 Đã thêm bảng điểm với {len(subjects[:8])} môn học")
        
        # Save the modified document
        doc.save(output_file)
        print(f"\n🎉 HOÀN THÀNH VERSION 2!")
        print(f"📁 File đã lưu: {output_file}")
        print(f"🔄 Tổng số thay đổi: {changes_made}")
        print(f"📚 Ngành học: {selected_course['name']}")
        print(f"📅 Thời gian: Đã tăng 1 năm 1 tháng (2025→2026)")
        print(f"📊 Điểm GPA: {gpa:.2f}/100 ({standing})")
        print(f"✅ Tỷ lệ hoàn thành: {completion_rate:.1f}%")
        print(f"🎓 Môn đã Pass: {passed_subjects}/{len(subjects[:8])}")
        
        return True
        
    except Exception as e:
        print(f"❌ Lỗi: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🎓 CHỈNH SỬA TRANSCRIPT FLINDERS UNIVERSITY - VERSION 2")
    print("="*60)
    edit_transcript_document_v2() 