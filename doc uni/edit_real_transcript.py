#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to edit REAL Flinders University transcript document
Based on actual transcript structure: Dick Jeffries, Education/Science degrees
Will update to IT/Engineering fields with +1 year +1 month timeline
"""

import docx
from docx import Document
import re
import random
import os

def edit_real_transcript():
    """Edit the real transcript document with realistic updates"""
    
    # File paths
    input_file = "test read doc/2386787-official-transcript-flinders-university.docx"
    output_file = "test read doc/Flinders_Real_Transcript_IT_Engineering_2023.docx"
    
    try:
        print("🔄 Đang đọc file transcript thực...")
        doc = Document(input_file)
        print(f"✅ Đã load document với {len(doc.paragraphs)} paragraphs")
        
        # New degree combinations (realistic IT/Engineering)
        degree_options = [
            {
                "degrees": "Bachelor of Engineering (Software Engineering) and Bachelor of Information Technology",
                "honors": "Bachelor of Engineering (Honours)",
                "teaching_areas": "Software Development\nDatabase Systems",
                "major": "Software Engineering"
            },
            {
                "degrees": "Bachelor of Science (Computer Science) and Bachelor of Information Technology", 
                "honors": "Bachelor of Science (Honours)",
                "teaching_areas": "Computer Science\nCybersecurity",
                "major": "Computer Science"
            },
            {
                "degrees": "Bachelor of Engineering (Computer Systems) and Bachelor of Science",
                "honors": "Bachelor of Engineering (Honours)", 
                "teaching_areas": "Network Engineering\nSystems Design",
                "major": "Computer Systems"
            }
        ]
        
        selected_degree = random.choice(degree_options)
        print(f"📚 Bằng cấp mới: {selected_degree['degrees']}")
        
        # Time updates (+1 year +1 month)
        year_mapping = {
            "2015": "2016", "2016": "2017", "2017": "2018", "2018": "2019",
            "2019": "2020", "2020": "2021", "2021": "2022", "2022": "2023"
        }
        
        # Month update (Apr 2022 -> May 2023)
        date_mapping = {
            "25 Apr 2022": "25 May 2023",
            "28 Jan 2021": "28 Feb 2022", 
            "25 Mar 2022": "25 Apr 2023",
            "17 Dec 2020": "17 Jan 2022"
        }
        
        # Subject mappings for IT/Engineering
        subject_replacements = {
            # Chemistry subjects -> Programming
            "CHEM1101": "COMP1101",
            "Chemical Structure and Bonding": "Introduction to Programming",
            "CHEM1102": "COMP1102", 
            "Modern Chemistry": "Software Development Fundamentals",
            "CHEM1202": "COMP1202",
            "Chemistry for Life Sciences": "Object-Oriented Programming",
            "CHEM2702": "COMP2702",
            "Organic Reactions": "Web Development",
            "CHEM2712": "COMP2712",
            "Separation Science": "Database Systems",
            "CHEM3702": "COMP3702",
            "Inorganic and Organometallic Chemistry": "Advanced Algorithms",
            "CHEM3711": "COMP3711",
            "Organic Synthesis and Mechanism": "Software Engineering",
            
            # Biology subjects -> Systems/Networks
            "BIOL1101": "ENGN1101",
            "Evolution of Biological Diversity": "Digital Systems Design",
            "BIOL1102": "ENGN1102",
            "Molecular Basis of Life": "Computer Systems Architecture", 
            "BIOL2712": "ENGN2712",
            "Animal Diversity": "Network Protocols",
            "BIOL2772": "ENGN2772",
            "Molecular Biology": "Operating Systems",
            "BIOL3771": "ENGN3771",
            "DNA to Genome": "Cloud Computing",
            "BIOL3731": "ENGN3731",
            "Aquaculture Reproduction": "Cybersecurity Fundamentals",
            
            # Education subjects -> IT Management
            "EDUC1120": "MGMT1120",
            "Teaching and Educational Contexts": "IT Project Management",
            "EDUC1223": "MGMT1223", 
            "Middle Schooling Philosophy and Pedagogy": "Systems Analysis and Design",
            "EDUC1227": "MGMT1227",
            "Professional Experience: Year 1": "Industry Placement: Year 1",
            "EDUC2320": "MGMT2320",
            "Learners and their Development": "Human-Computer Interaction",
            "EDUC2326": "MGMT2326",
            "Professional Experience: Year 2A": "Industry Placement: Year 2A",
            "EDUC2420": "MGMT2420", 
            "Teaching Indigenous Australian Students": "Ethics in Technology",
            "EDUC2426": "MGMT2426",
            "Professional Experience: Year 2B": "Industry Placement: Year 2B",
            
            # More Education -> IT subjects
            "EDUC3526": "COMP3526",
            "Literacies Across the Curriculum": "Mobile Application Development",
            "EDUC3530": "COMP3530",
            "Professional Experience: Year 3A": "Industry Project: Year 3A", 
            "EDUC3531H": "COMP3531H",
            "Science Curriculum Specialisation A1": "Advanced Programming A1",
            "EDUC3620": "COMP3620",
            "Relationships for Learning": "User Experience Design",
            "EDUC3627": "COMP3627",
            "Contemporary Issues in Secondary Schooling": "Contemporary Issues in Technology",
            "EDUC3642": "COMP3642",
            "Professional Experience: Year 3B": "Industry Project: Year 3B",
            "EDUC3643H": "COMP3643H",
            "Science Curriculum Specialisation B1": "Advanced Programming B1",
            "EDUC3628": "COMP3628",
            "Numeracy and ICT Across the Middle and Secondary Curriculum": "Data Analytics and Visualization",
            
            # Final year subjects
            "EDUC4720": "COMP4720",
            "Differentiation for Diverse Learners": "Machine Learning Applications",
            "EDUC4729K": "ENGN4729K",
            "Chemistry, Earth Sciences, Physics Curriculum Specialisation": "Network Security and Cryptography",
            "EDUC4730M": "COMP4730M", 
            "Biological Sciences Curriculum Specialisation": "Artificial Intelligence Systems",
            "EDUC4742": "COMP4742",
            "Professional Experience: Final Assessment": "Capstone Project: Final Assessment",
            "EDUC4820": "MGMT4820",
            "The Professional Educator": "Professional IT Practice",
            "EDUC4744": "COMP4744",
            "Extended Urban Professional Experience": "Extended Industry Experience",
            
            # Honours subjects
            "EDUC7120": "COMP7120",
            "Approaches to Research (Honours)": "Research Methods in Computer Science",
            "EDUC7121": "COMP7121", 
            "Honours Thesis Preparation": "Honours Project Preparation",
            "EDUC7122": "COMP7122",
            "Honours Research Methods": "Advanced Research Methods",
            "EDUC7220": "COMP7220",
            "Honours Thesis": "Honours Research Project",
            
            # Other subjects
            "FACH1701": "COMP1701",
            "Introduction to Forensic Science": "Computer Forensics",
            "NANO2701": "ENGN2701",
            "Structure and Characterisation": "Hardware Design"
        }
        
        # Process document content
        changes_made = 0
        
        print("🔄 Đang xử lý nội dung document...")
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            new_text = original_text
            
            # Update student info and degrees
            if "Bachelor of Education (Middle and Secondary Schooling) and Bachelor of Science" in new_text:
                new_text = new_text.replace(
                    "Bachelor of Education (Middle and Secondary Schooling) and Bachelor of Science",
                    selected_degree['degrees']
                )
                changes_made += 1
            
            if "Bachelor of Education (Honours)" in new_text:
                new_text = new_text.replace("Bachelor of Education (Honours)", selected_degree['honors'])
                changes_made += 1
            
            if "Honours Degree of Bachelor of Education" in new_text:
                new_text = new_text.replace("Honours Degree of Bachelor of Education", f"Honours Degree of {selected_degree['honors'].split(' (')[0]}")
                changes_made += 1
            
            # Update teaching areas
            if "Teaching Area:  Biology\nChemistry" in new_text:
                new_text = new_text.replace("Teaching Area:  Biology\nChemistry", f"Specialisation:  {selected_degree['teaching_areas']}")
                changes_made += 1
            
            # Update major
            if "Major:	Mathematics" in new_text:
                new_text = new_text.replace("Major:	Mathematics", f"Major:	{selected_degree['major']}")
                changes_made += 1
            
            # Update all years
            for old_year, new_year in year_mapping.items():
                if old_year in new_text:
                    new_text = new_text.replace(old_year, new_year)
                    changes_made += 1
            
            # Update specific dates
            for old_date, new_date in date_mapping.items():
                if old_date in new_text:
                    new_text = new_text.replace(old_date, new_date)
                    changes_made += 1
            
            # Update subject codes and names
            for old_subject, new_subject in subject_replacements.items():
                if old_subject in new_text:
                    new_text = new_text.replace(old_subject, new_subject)
                    changes_made += 1
            
            # Apply changes
            if new_text != original_text:
                paragraph.text = new_text
                if len(original_text) > 20:  # Only show substantial changes
                    print(f"✏️  Sửa: {original_text[:50]}... → {new_text[:50]}...")
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = original_text
                    
                    # Apply same replacements to table content
                    for old_year, new_year in year_mapping.items():
                        if old_year in new_text:
                            new_text = new_text.replace(old_year, new_year)
                            changes_made += 1
                    
                    for old_date, new_date in date_mapping.items():
                        if old_date in new_text:
                            new_text = new_text.replace(old_date, new_date)
                            changes_made += 1
                    
                    for old_subject, new_subject in subject_replacements.items():
                        if old_subject in new_text:
                            new_text = new_text.replace(old_subject, new_subject)
                            changes_made += 1
                    
                    if new_text != original_text:
                        cell.text = new_text
        
        # Save document
        doc.save(output_file)
        
        print(f"\n🎉 HOÀN THÀNH CHỈNH SỬA TRANSCRIPT THỰC!")
        print(f"📁 File mới: {output_file}")
        print(f"🔄 Tổng thay đổi: {changes_made}")
        print(f"📚 Bằng cấp: {selected_degree['degrees']}")
        print(f"🎓 Honours: {selected_degree['honors']}")
        print(f"📅 Thời gian: Tăng 1 năm 1 tháng (2015-2022 → 2016-2023)")
        print(f"📋 Chuyên ngành: {selected_degree['major']}")
        print(f"💼 Lĩnh vực: {selected_degree['teaching_areas'].replace(chr(10), ', ')}")
        
        return True
        
    except Exception as e:
        print(f"❌ Lỗi: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🎓 CHỈNH SỬA TRANSCRIPT THỰC - FLINDERS UNIVERSITY")
    print("="*60)
    print("👤 Sinh viên: Dick Jeffries → Chuyển sang IT/Engineering")
    print("📅 Thời gian: 2015-2022 → 2016-2023")
    print("📚 Ngành: Education/Science → IT/Engineering")
    print("="*60)
    edit_real_transcript() 